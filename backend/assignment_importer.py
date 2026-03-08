import io
import pandas as pd
from pydantic import BaseModel
from typing import List, Optional
from thefuzz import process, fuzz
import re

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

class ExtractedAssignment(BaseModel):
    raw_teacher: str
    raw_subject: str
    theory_credits: int = 3
    lab_credits: int = 0
    raw_lab_engineer: Optional[str] = None
    raw_batch: Optional[str] = None
    raw_sections: Optional[str] = None  # e.g. "A+B+C"
    
    # Matched data (to be filled by backend)
    matched_teacher_id: Optional[int] = None
    matched_teacher_name: Optional[str] = None
    matched_subject_id: Optional[int] = None
    matched_subject_name: Optional[str] = None
    matched_lab_engineer_id: Optional[int] = None
    matched_lab_engineer_name: Optional[str] = None
    matched_batch_id: Optional[int] = None
    matched_batch_name: Optional[str] = None
    confidence: float = 0.0

def clean_text(text: str) -> str:
    if not text: return ""
    return re.sub(r'\s+', ' ', text).strip()

def extract_from_pdf(content: bytes) -> str:
    if not PdfReader: return ""
    reader = PdfReader(io.BytesIO(content))
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_from_docx(content: bytes) -> str:
    if not Document: return ""
    doc = Document(io.BytesIO(content))
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            text += " | ".join([cell.text for cell in row.cells]) + "\n"
    return text

def extract_docx_tables(content: bytes) -> pd.DataFrame:
    """Extract tables from DOCX as a DataFrame for structured parsing."""
    if not Document:
        return None
    doc = Document(io.BytesIO(content))
    all_rows = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            all_rows.append(cells)
    if not all_rows:
        return None
    # Use first row as headers if it looks like a header
    max_cols = max(len(r) for r in all_rows)
    # Pad shorter rows
    for r in all_rows:
        while len(r) < max_cols:
            r.append('')
    
    # Check if first row looks like a header (contains mostly text, not numbers)
    first_row = all_rows[0]
    is_header = sum(1 for c in first_row if c and not re.match(r'^\d+\.?\d*$', c.strip())) >= len(first_row) * 0.5
    
    if is_header and len(all_rows) > 1:
        df = pd.DataFrame(all_rows[1:], columns=first_row)
    else:
        df = pd.DataFrame(all_rows)
    return df


# ── Helpers ────────────────────────────────────────────────────

# Salutation prefixes that indicate a teacher name
TEACHER_PREFIXES = re.compile(r'(Prof\.?|Dr\.?|Engr\.?|Mr\.?|Ms\.?|Mrs\.?|Sir|Madam)\s', re.IGNORECASE)

# Common batch/program patterns
BATCH_PATTERN = re.compile(
    r'\b(B\.?\s*(?:CS|SE|EE|CE|AE|AI|ME|IT|CY|DS|BBA|BA)|M\.?\s*(?:CS|SE|EE|CE)|BS\w*|MS\w*|PhD)[-\s]*(\d{1,2})?\s*([A-Z])?\b',
    re.IGNORECASE
)
BATCH_YEAR_PATTERN = re.compile(r'\b(?:Batch|Session|Year)\s*(\d{4})\b', re.IGNORECASE)

# Section pattern in brackets: (A+B+C), (A, B, C), (Section A+B)
SECTION_PATTERN = re.compile(r'\(([^)]*(?:[A-Z]\s*[\+\,\&]\s*[A-Z])[^)]*)\)', re.IGNORECASE)
SECTION_PATTERN_ALT = re.compile(r'\(\s*([A-Z](?:\s*[\+\,\&\s]\s*[A-Z])*)\s*(?:section)?\s*\)', re.IGNORECASE)

# Credit pattern: (3+0), (2+1), 3+1, 3-0
CREDIT_PATTERN = re.compile(r'\(?\s*(\d)\s*[\+\-\,]\s*(\d)\s*\)?')


def is_valid_name(val) -> bool:
    """Check if a value looks like a real person/subject name, not a number or junk."""
    if val is None:
        return False
    s = str(val).strip()
    if not s or len(s) < 2:
        return False
    # Reject pure numbers, floats, NaN
    if re.match(r'^[\d\.\-\+\s]+$', s):
        return False
    if s.lower() in ('nan', 'none', 'null', ''):
        return False
    # Reject if it's just special chars
    if re.match(r'^[\W_]+$', s):
        return False
    return True


def extract_batch_from_line(line):
    """Try to extract a batch/program reference from a line."""
    m = BATCH_PATTERN.search(line)
    if m:
        return m.group(0).strip()
    m = BATCH_YEAR_PATTERN.search(line)
    if m:
        return m.group(0).strip()
    return None


def extract_sections_from_text(text):
    """Extract section letters from text like 'Dr. Ali (A+B+C section)'."""
    m = SECTION_PATTERN.search(text)
    if not m:
        m = SECTION_PATTERN_ALT.search(text)
    if m:
        return m.group(1).strip()
    return None


def strip_sections_from_name(name):
    """Remove section info in brackets from a teacher name."""
    cleaned = re.sub(r'\s*\([^)]*\)\s*', ' ', name).strip()
    return cleaned


def identify_teacher_part(parts):
    """Given list of text parts, identify which is most likely the teacher name."""
    for i, p in enumerate(parts):
        if TEACHER_PREFIXES.search(p):
            return i
    return 0


# ── Column-based parser for Excel / CSV / DOCX tables ──────────

def _find_col(columns, keywords):
    """Find a column whose header fuzzy-matches any of the given keywords."""
    cols_lower = {c: str(c).lower().strip() for c in columns}
    for kw in keywords:
        kw_l = kw.lower()
        for col, col_lower in cols_lower.items():
            if kw_l in col_lower or col_lower in kw_l:
                return col
    return None


def _ensure_scalar(val):
    """Ensure a value is a scalar, not a Pandas Series or other collection."""
    if isinstance(val, (pd.Series, list, tuple)):
        return val[0] if len(val) > 0 else None
    return val


def parse_dataframe(df) -> List[dict]:
    """
    Smart column-based parser for Excel/CSV/DOCX-table DataFrames.
    Identifies columns by header names and extracts assignment rows.
    """
    if df is None or df.empty:
        return []
    
    # Deduplicate column names before anything else. 
    # Duplicate headers cause row.get(col) to return a Series, crashing bool checks.
    cols = []
    counts = {}
    for c in df.columns:
        c_str = str(c).strip()
        if c_str in counts:
            counts[c_str] += 1
            cols.append(f"{c_str}_{counts[c_str]}")
        else:
            counts[c_str] = 0
            cols.append(c_str)
    df.columns = cols
    
    # Try to identify columns by common header keywords
    teacher_col = _find_col(df.columns, [
        'teacher', 'faculty', 'instructor', 'professor', 'name of teacher',
        'faculty name', 'teacher name', 'name of faculty', 'assigned to', 'taught by'
    ])
    subject_col = _find_col(df.columns, [
        'subject', 'course', 'course title', 'subject name', 'course name',
        'module', 'paper'
    ])
    code_col = _find_col(df.columns, [
        'course code', 'subject code', 'course no', 'course id', 'code'
    ])
    credits_col = _find_col(df.columns, [
        'credit hrs', 'credit hours', 'credit', 'cr', 'ch', 'credits'
    ])
    theory_col = _find_col(df.columns, [
        'theory', 'theory credits', 'theory cr', 'lec', 'lecture'
    ])
    lab_col_header = _find_col(df.columns, [
        'lab credits', 'lab cr', 'practical', 'lab hours'
    ])
    batch_col = _find_col(df.columns, [
        'batch', 'program', 'class', 'semester', 'degree'
    ])
    lab_eng_col = _find_col(df.columns, [
        'lab engineer', 'lab engr', 'lab instructor', 'lab assistant', 'lab staff'
    ])
    designation_col = _find_col(df.columns, [
        'designation', 'rank', 'title', 'position'
    ])
    
    # If we couldn't find teacher or subject columns, try heuristic:
    # skip numeric columns and pick text-heavy columns
    if teacher_col is None or subject_col is None:
        text_cols = []
        for col in df.columns:
            sample = df[col].dropna().head(10)
            name_count = sum(1 for v in sample if is_valid_name(v))
            if name_count >= max(1, len(sample) * 0.3):
                text_cols.append(col)
        
        if len(text_cols) >= 2:
            if not teacher_col:
                # Pick the column where more values have salutation prefixes
                teacher_scores = []
                for col in text_cols:
                    salutation_count = sum(
                        1 for v in df[col].dropna().head(20)
                        if TEACHER_PREFIXES.search(str(v).strip())
                    )
                    teacher_scores.append((col, salutation_count))
                teacher_scores.sort(key=lambda x: x[1], reverse=True)
                teacher_col = teacher_scores[0][0]
            
            if not subject_col:
                subject_col = next((c for c in text_cols if c != teacher_col), None)
        elif len(text_cols) == 1 and not teacher_col:
            teacher_col = text_cols[0]
    
    if not teacher_col and not subject_col:
        return []
    
    rows = []
    for _, row in df.iterrows():
        raw_teacher = _ensure_scalar(row.get(teacher_col, ''))
        raw_teacher = str(raw_teacher).strip() if raw_teacher is not None else ''
        
        raw_subject = _ensure_scalar(row.get(subject_col, ''))
        raw_subject = str(raw_subject).strip() if raw_subject is not None else ''
        
        # Skip rows where both teacher and subject look invalid
        if not is_valid_name(raw_teacher) and not is_valid_name(raw_subject):
            continue
        
        # Extract sections from teacher name like "Dr. Ali (A+B+C section)"
        raw_sections = extract_sections_from_text(raw_teacher) if is_valid_name(raw_teacher) else None
        if raw_sections:
            raw_teacher = strip_sections_from_name(raw_teacher)
        
        # If teacher isn't valid but subject is, swap or skip
        if not is_valid_name(raw_teacher) and is_valid_name(raw_subject):
            raw_teacher = ''
        if not is_valid_name(raw_subject):
            raw_subject = ''
        
        # Parse credits
        theory = 3
        lab = 0
        
        if theory_col is not None and pd.notna(row.get(theory_col)):
            try:
                theory_val = _ensure_scalar(row[theory_col])
                theory = int(float(str(theory_val).strip()))
            except (ValueError, TypeError):
                pass
        if lab_col_header is not None and pd.notna(row.get(lab_col_header)):
            try:
                lab_val = _ensure_scalar(row[lab_col_header])
                lab = int(float(str(lab_val).strip()))
            except (ValueError, TypeError):
                pass
        
        # Combined credit column (e.g. "3+1", "(3+0)")
        if credits_col is not None and pd.notna(row.get(credits_col)):
            cr_val = _ensure_scalar(row[credits_col])
            header_val = str(cr_val).strip()
            cr_match = CREDIT_PATTERN.search(header_val)
            if cr_match:
                theory = int(cr_match.group(1))
                lab = int(cr_match.group(2))
            elif theory_col is None:
                try:
                    theory = int(float(header_val))
                except (ValueError, TypeError):
                    pass
        
        # Batch info
        raw_batch = None
        if batch_col is not None and pd.notna(row.get(batch_col)):
            batch_val = _ensure_scalar(row[batch_col])
            batch_str = str(batch_val).strip()
            if is_valid_name(batch_str):
                raw_batch = batch_str
        if not raw_batch:
            row_text = ' '.join(str(v) for v in row.values if pd.notna(v))
            raw_batch = extract_batch_from_line(row_text)
        
        # Lab engineer
        raw_lab_engineer = None
        if lab_eng_col is not None and pd.notna(row.get(lab_eng_col)):
            le_val = _ensure_scalar(row[lab_eng_col])
            le_str = str(le_val).strip()
            if is_valid_name(le_str):
                raw_lab_engineer = le_str
        
        # Subject code prefix
        if code_col is not None and pd.notna(row.get(code_col)) and is_valid_name(raw_subject):
            code_val = _ensure_scalar(row[code_col])
            code_str = str(code_val).strip()
            if code_str and code_str not in raw_subject:
                raw_subject = f"{code_str} - {raw_subject}"
        
        rows.append({
            "raw_teacher": raw_teacher or 'Unknown',
            "raw_subject": raw_subject or 'Unknown',
            "theory_credits": theory,
            "lab_credits": lab,
            "raw_batch": raw_batch,
            "raw_lab_engineer": raw_lab_engineer,
            "raw_sections": raw_sections,
        })
    
    return rows


# ── Text-based parser for PDF / plain text ──────────────────────

def parse_text_to_rows(text: str) -> List[dict]:
    """
    Heuristic parser for text extracted from PDF/DOCX.
    
    Known Word table format (pipe-delimited after extraction):
        S.No | Subject Name | 3+0 | Teacher Name (A+B+C) | Lab Engineer
    
    The teacher name comes AFTER the credit hours, with optional section info in brackets.
    """
    rows = []
    lines = text.split("\n")
    
    for line in lines:
        line = clean_text(line)
        if not line or len(line) < 8:
            continue
        
        credit_match = CREDIT_PATTERN.search(line)
        if credit_match:
            theory = int(credit_match.group(1))
            lab = int(credit_match.group(2))
            
            # Split the ENTIRE line by pipes/tabs
            parts = [p.strip() for p in re.split(r'[|\t]', line) if p.strip()]
            
            if len(parts) >= 3:
                # Filter: find which parts are valid names vs serial numbers vs credits
                valid_parts = []
                credit_part_idx = None
                
                for i, p in enumerate(parts):
                    # Check if this part IS the credit info
                    if CREDIT_PATTERN.search(p) and credit_part_idx is None:
                        credit_part_idx = i
                    elif is_valid_name(p):
                        valid_parts.append((i, p))
                
                if len(valid_parts) < 1:
                    continue
                
                # Known Word format: S.No | Subject | Credits | Teacher (sections) | Lab Engr
                # After filtering: valid_parts has Subject, Teacher, maybe Lab Engr
                
                raw_subject = ''
                raw_teacher = ''
                raw_lab_engineer = None
                raw_sections = None
                raw_batch = extract_batch_from_line(line)
                
                if credit_part_idx is not None:
                    # Parts BEFORE credit = subject(s), parts AFTER credit = teacher + lab engr
                    before_credit = [(i, p) for i, p in valid_parts if i < credit_part_idx]
                    after_credit = [(i, p) for i, p in valid_parts if i > credit_part_idx]
                    
                    if before_credit:
                        raw_subject = before_credit[-1][1]  # Last valid part before credit
                    
                    if after_credit:
                        # First part after credit = teacher (may have sections in brackets)
                        teacher_text = after_credit[0][1]
                        raw_sections = extract_sections_from_text(teacher_text)
                        raw_teacher = strip_sections_from_name(teacher_text) if raw_sections else teacher_text
                        
                        # If there are more parts after credit, could be lab engineer
                        if len(after_credit) > 1:
                            raw_lab_engineer = after_credit[1][1]
                    
                    # If no teacher found after credit, try salutation-based detection
                    if not raw_teacher and before_credit:
                        for _, p in before_credit:
                            if TEACHER_PREFIXES.search(p):
                                raw_teacher = p
                                break
                else:
                    # No clear credit position - use salutation heuristic
                    teacher_idx = None
                    for vi, (_, p) in enumerate(valid_parts):
                        if TEACHER_PREFIXES.search(p):
                            teacher_idx = vi
                            break
                    
                    if teacher_idx is not None:
                        raw_teacher = valid_parts[teacher_idx][1]
                        raw_sections = extract_sections_from_text(raw_teacher)
                        if raw_sections:
                            raw_teacher = strip_sections_from_name(raw_teacher)
                        remaining = [p for vi, (_, p) in enumerate(valid_parts) if vi != teacher_idx]
                        if remaining:
                            raw_subject = remaining[0]
                        if len(remaining) > 1:
                            raw_lab_engineer = remaining[1]
                    elif len(valid_parts) >= 2:
                        raw_teacher = valid_parts[0][1]
                        raw_subject = valid_parts[1][1]
                        if len(valid_parts) > 2:
                            raw_lab_engineer = valid_parts[2][1]
                
                if raw_teacher or raw_subject:
                    rows.append({
                        "raw_teacher": raw_teacher or 'Unknown',
                        "raw_subject": raw_subject or 'Unknown',
                        "theory_credits": theory,
                        "lab_credits": lab,
                        "raw_batch": raw_batch,
                        "raw_lab_engineer": raw_lab_engineer,
                        "raw_sections": raw_sections,
                    })
                    
    return rows


# ── Fuzzy matcher ──────────────────────────────────────────────

def fuzzy_match_entity(raw_name: str, choices: List[dict], key="name", threshold=60):
    """
    Fuzzy match a raw name against a list of DB entities.
    choices: list of {'id': 1, 'name': 'Full Name'}
    """
    if not raw_name or not choices: return None, 0
    
    names = [c[key] for c in choices]
    best_match, score = process.extractOne(raw_name, names, scorer=fuzz.token_sort_ratio)
    
    if score >= threshold:
        matched_item = next(c for c in choices if c[key] == best_match)
        return matched_item, score
    return None, score
